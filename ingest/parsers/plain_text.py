"""Faithful, lossless document text extraction.

The RAG chunkers (`parse_docx` / `parse_pdf`) split documents into clause-level
chunks with section detection and noise-section stripping — the right behavior
for retrieval. They are NOT safe as plain text extractors: a contract's clause
body packed into a numbered paragraph is mis-read as a heading and discarded
(see the audit; ~89% coverage on a real NDA, losing the exclusions /
indemnification / execution clauses).

The contract-review upload path needs the VERBATIM document — every paragraph
and table, in order — the same fidelity the Word add-in gets from Office.js
`body.getText()`. These extractors provide exactly that, with no heuristics.
"""

from __future__ import annotations

from pathlib import Path

import pdfplumber
from docx import Document
from docx.oxml.ns import qn


def extract_docx_text(filepath: Path | str) -> str:
    """Return every body paragraph + table cell, in document order. Lossless.

    Iterates the body so paragraphs and tables stay interleaved in place. No
    section detection, no noise filtering, no chunk splitting — that is the
    chunker's job, not this one.
    """
    doc = Document(str(filepath))
    para_map = {p._element: p for p in doc.paragraphs}
    tbl_map = {t._element: t for t in doc.tables}

    parts: list[str] = []
    for element in doc.element.body:
        if element.tag == qn("w:p"):
            para = para_map.get(element)
            if para is not None and para.text.strip():
                parts.append(para.text)
        elif element.tag == qn("w:tbl"):
            table = tbl_map.get(element)
            if table is not None:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append("\t".join(cells))
    return "\n".join(parts)


def extract_pdf_text(filepath: Path | str) -> str:
    """Return the text of every page, in order, joined by blank lines."""
    parts: list[str] = []
    with pdfplumber.open(str(filepath)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
    return "\n\n".join(parts)


def extract_document_text(filepath: Path | str, suffix: str | None = None) -> str:
    """Dispatch to the right faithful extractor by file extension.

    `suffix` overrides the on-disk extension — the upload path hands a temp file
    whose suffix may be wrong or absent, but knows the original filename's
    extension. Raises ValueError for unsupported types.
    """
    filepath = Path(filepath)
    ext = (suffix or filepath.suffix).lower()
    if ext == ".docx":
        return extract_docx_text(filepath)
    if ext == ".pdf":
        return extract_pdf_text(filepath)
    if ext == ".txt":
        return filepath.read_text(encoding="utf-8")
    raise ValueError(f"Unsupported file type: {ext!r} (supported: .docx, .pdf, .txt)")
