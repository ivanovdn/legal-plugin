# tests/test_plain_text.py
"""Faithful document text extraction for the contract-review / upload path.

Unlike the RAG chunker (`parse_docx` / `parse_pdf`), these extractors must be
LOSSLESS: the contract review feeds the whole document to the LLM, so every
paragraph and table cell must survive. The chunker is a clause-splitter with
section-detection + noise-stripping heuristics that silently drop content when
(mis)used as a text extractor — see docs/wiki.md and the audit notes.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from ingest.parsers.plain_text import (
    extract_docx_text,
    extract_document_text,
)

ROOT = Path(__file__).resolve().parent.parent
NDA = ROOT / "data" / "Trinetix Model NDA 2025.docx"


def test_extract_docx_text_keeps_every_paragraph_and_cell(tmp_path):
    """Every body paragraph and table cell must appear in the output."""
    doc = Document()
    doc.add_paragraph("Recital paragraph one.")
    # A long, numbered, inline clause body — the exact shape the RAG chunker
    # mis-classifies as a heading and discards. The faithful extractor must keep it.
    doc.add_paragraph(
        "1. The Receiving Party shall be liable for any unauthorized disclosure "
        "and shall indemnify the Disclosing Party against all resulting losses."
    )
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Signed by:"
    table.rows[0].cells[1].text = "Boris"
    table.rows[1].cells[0].text = "Title:"
    table.rows[1].cells[1].text = "CEO"
    doc.add_paragraph("Final execution paragraph.")
    fp = tmp_path / "sample.docx"
    doc.save(str(fp))

    out = extract_docx_text(fp)

    for para in ("Recital paragraph one.", "shall indemnify the Disclosing Party",
                 "Final execution paragraph."):
        assert para in out, f"paragraph text dropped: {para!r}"
    for cell in ("Signed by:", "Boris", "Title:", "CEO"):
        assert cell in out, f"table cell dropped: {cell!r}"


def test_extract_docx_text_preserves_document_order(tmp_path):
    """A table in the middle of the doc appears in place, not appended at the end."""
    doc = Document()
    doc.add_paragraph("ALPHA before table")
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "MIDDLE cell"
    doc.add_paragraph("OMEGA after table")
    fp = tmp_path / "ordered.docx"
    doc.save(str(fp))

    out = extract_docx_text(fp)
    assert out.index("ALPHA") < out.index("MIDDLE") < out.index("OMEGA")


@pytest.mark.skipif(not NDA.exists(), reason="real contract fixture is gitignored / not present")
def test_extract_docx_recovers_clauses_the_chunker_drops():
    """Regression on the real NDA: the faithful extractor keeps the three clauses
    that `parse_docx` (as an extractor) silently dropped. See the audit."""
    out = extract_docx_text(NDA).lower()
    for phrase in ("this definition excludes", "indemnify", "in witness whereof"):
        assert phrase in out, f"faithful extractor missing {phrase!r}"

    # Contrast: the RAG chunker, rejoined as text, drops the execution block.
    from ingest.parsers.docx_parser import parse_docx
    chunked = "\n\n".join(
        c.text for c in parse_docx(
            NDA, client_id="internal", jurisdiction="", doc_type="contract", sensitivity="internal"
        )
    ).lower()
    assert "in witness whereof" not in chunked, "chunker unexpectedly kept the execution block"


def test_extract_document_text_reads_txt(tmp_path):
    fp = tmp_path / "note.txt"
    fp.write_text("plain contract text", encoding="utf-8")
    assert extract_document_text(fp) == "plain contract text"


def test_extract_document_text_suffix_override_routes_docx(tmp_path):
    """Chainlit hands a temp path whose suffix may be wrong/absent; the explicit
    suffix hint must drive the routing."""
    doc = Document()
    doc.add_paragraph("Body clause kept.")
    fp = tmp_path / "upload.bin"  # wrong suffix on disk
    doc.save(str(fp))
    out = extract_document_text(fp, suffix=".docx")
    assert "Body clause kept." in out


def test_extract_document_text_routes_pdf(tmp_path, monkeypatch):
    import ingest.parsers.plain_text as pt
    monkeypatch.setattr(pt, "extract_pdf_text", lambda fp: "PDF-TEXT")
    fp = tmp_path / "doc.pdf"
    fp.write_bytes(b"%PDF-1.4\n")
    assert pt.extract_document_text(fp) == "PDF-TEXT"


def test_extract_document_text_unsupported_raises(tmp_path):
    fp = tmp_path / "thing.xyz"
    fp.write_text("x", encoding="utf-8")
    with pytest.raises(ValueError):
        extract_document_text(fp)
