# tests/test_parsers.py
from pathlib import Path
from ingest.chunk_models import LegalChunk


def test_docx_parser_returns_legal_chunks(tmp_path):
    """DOCX parser returns list of LegalChunk objects."""
    from docx import Document

    doc = Document()
    doc.add_heading("Section 1: Definitions", level=1)
    doc.add_paragraph("This agreement defines the terms used herein.")
    doc.add_heading("Section 2: Obligations", level=1)
    doc.add_paragraph("The contractor shall deliver services as described.")
    filepath = tmp_path / "test_contract.docx"
    doc.save(str(filepath))

    from ingest.parsers.docx_parser import parse_docx

    chunks = parse_docx(
        filepath=filepath,
        client_id="client-abc",
        jurisdiction="US-DE",
        doc_type="contract",
        sensitivity="confidential",
    )

    assert len(chunks) > 0
    assert all(isinstance(c, LegalChunk) for c in chunks)
    assert all(c.client_id == "client-abc" for c in chunks)
    assert all(c.jurisdiction == "US-DE" for c in chunks)
    assert all(c.doc_type == "contract" for c in chunks)
    assert all(c.sensitivity == "confidential" for c in chunks)
    assert all(c.doc_filename == "test_contract.docx" for c in chunks)


def test_docx_parser_chunk_has_text(tmp_path):
    """Each chunk has non-empty text."""
    from docx import Document

    doc = Document()
    doc.add_heading("Important Clause", level=1)
    doc.add_paragraph("The parties hereby agree to binding arbitration for all disputes arising under this agreement.")
    filepath = tmp_path / "test.docx"
    doc.save(str(filepath))

    from ingest.parsers.docx_parser import parse_docx

    chunks = parse_docx(
        filepath=filepath,
        client_id="internal",
        jurisdiction="US",
        doc_type="policy",
        sensitivity="internal",
    )

    assert all(len(c.text.strip()) > 0 for c in chunks)


def test_pdf_parser_returns_legal_chunks():
    """PDF parser returns list of LegalChunk objects from text."""
    from ingest.parsers.pdf_parser import parse_pdf_text

    text = """Section 1: Definitions
This agreement defines the terms used herein and throughout the document for all parties involved in this contractual relationship. The term Service Provider refers to the entity providing professional services under the terms and conditions set forth in this agreement. The term Client refers to the entity receiving and paying for such professional services as described in the attached schedules and exhibits.

Section 2: Obligations
The contractor shall deliver services as described in Exhibit A attached hereto and incorporated by reference into this agreement. All deliverables must meet the quality standards specified in the service level agreement. The contractor agrees to maintain adequate staffing levels and professional certifications as required by applicable law and industry best practices throughout the term of this agreement."""

    chunks = parse_pdf_text(
        text=text,
        filename="test.pdf",
        client_id="client-y",
        jurisdiction="UK",
        doc_type="legislation",
        sensitivity="public",
    )

    assert len(chunks) > 0
    assert all(isinstance(c, LegalChunk) for c in chunks)
    assert all(c.client_id == "client-y" for c in chunks)
    assert all(c.doc_type == "legislation" for c in chunks)
    assert all(c.doc_filename == "test.pdf" for c in chunks)
